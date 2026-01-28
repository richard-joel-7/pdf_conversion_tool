import os
import sys
import numpy as np
import cv2
from PIL import Image
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm
import glob
from bs4 import BeautifulSoup
import re

# Configuration
# ==============================================================================

# 1. Tesseract Executable Path
possible_tesseract_paths = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    os.path.join(os.getenv('LOCALAPPDATA'), r"Tesseract-OCR\tesseract.exe")
]
TESSERACT_CMD = None
for path in possible_tesseract_paths:
    if os.path.exists(path):
        TESSERACT_CMD = path
        break

if not TESSERACT_CMD:
    print("Warning: Tesseract executable not found in common locations.")
    print("Please ensure Tesseract is installed and added to PATH, or update the script.")
else:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# 2. Poppler Path
POPPLER_PATH = r"C:\Users\richardjoel.d\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe\poppler-25.07.0\Library\bin"

# 3. Tesseract Data Path
TESSDATA_CONFIG = ""
if TESSERACT_CMD:
    tess_base = os.path.dirname(TESSERACT_CMD)
    sys_tessdata = os.path.join(tess_base, "tessdata")
    if not os.path.exists(os.path.join(sys_tessdata, "tam.traineddata")):
        local_tessdata = os.path.join(os.getcwd(), "tessdata")
        if os.path.exists(os.path.join(local_tessdata, "tam.traineddata")):
            os.environ["TESSDATA_PREFIX"] = local_tessdata
            print(f"Using local tessdata: {local_tessdata}")

# ==============================================================================

def preprocess_image(pil_image):
    """
    Increases contrast and converts to grayscale to improve OCR accuracy.
    """
    open_cv_image = np.array(pil_image) 
    if len(open_cv_image.shape) == 3:
        open_cv_image = open_cv_image[:, :, ::-1].copy()

    gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(thresh)

def parse_bbox(title_str):
    """Parses 'bbox x0 y0 x1 y1' from title attribute"""
    if not title_str:
        return None
    match = re.search(r'bbox (\d+) (\d+) (\d+) (\d+)', title_str)
    if match:
        return [int(g) for g in match.groups()]
    return None

def hocr_to_docx(hocr_content, doc, page_num):
    """
    Parses HOCR content and adds it to the DOCX document with layout approximation.
    """
    soup = BeautifulSoup(hocr_content, 'lxml')
    
    # Get page dimensions if available
    page_div = soup.find('div', class_='ocr_page')
    page_width = 1000 # default
    if page_div:
        bbox = parse_bbox(page_div.get('title'))
        if bbox:
            page_width = bbox[2] - bbox[0]

    # Iterate paragraphs
    paragraphs = soup.find_all('p', class_='ocr_par')
    
    for p_idx, p in enumerate(paragraphs):
        # Create a new paragraph in DOCX
        docx_p = doc.add_paragraph()
        
        # Analyze paragraph alignment/indentation
        # We look at the first line's bounding box
        lines = p.find_all('span', class_='ocr_line')
        if not lines:
            continue
            
        first_line = lines[0]
        bbox = parse_bbox(first_line.get('title'))
        
        align = WD_ALIGN_PARAGRAPH.LEFT
        indent = 0
        
        if bbox:
            x0, y0, x1, y1 = bbox
            line_width = x1 - x0
            x_center = (x0 + x1) / 2
            
            # Center detection heuristic
            # If line center is within 10% of page center
            page_center = page_width / 2
            if abs(x_center - page_center) < (page_width * 0.1):
                align = WD_ALIGN_PARAGRAPH.CENTER
            # Right align heuristic (less common in scripts, but possible)
            elif x1 > (page_width * 0.9) and x0 > (page_width * 0.5):
                align = WD_ALIGN_PARAGRAPH.RIGHT
            # Indentation
            elif x0 > (page_width * 0.1):
                # Approximate indentation in inches (assuming 300 DPI or relative)
                # Word standard indent is usually 0.5 inch.
                # Let's map x0 percentage to inches.
                # A4 width is ~8.27 inches.
                indent_ratio = x0 / page_width
                indent_inches = indent_ratio * 8.27
                # Cap indent to avoid extreme values
                indent = min(indent_inches, 4.0)

        docx_p.alignment = align
        if indent > 0:
            docx_p.paragraph_format.left_indent = Inches(indent)

        # Iterate lines and words
        full_text = ""
        for line in lines:
            words = line.find_all('span', class_='ocrx_word')
            line_text = []
            
            # Font detection (Bold)
            # Tesseract usually doesn't output font style in standard HOCR unless configured.
            # We can try to detect based on 'strong' tag if present (rare) or text content.
            # For now, we'll stick to text content to ensure editability.
            
            for word in words:
                text = word.get_text().strip()
                if text:
                    # Check for bold confidence or style? 
                    # Tesseract doesn't give 'bold' flag easily.
                    # We will add text as normal run.
                    line_text.append(text)
            
            if line_text:
                full_text += " ".join(line_text) + " "
        
        # Clean up text
        full_text = full_text.strip()
        if full_text:
            run = docx_p.add_run(full_text)
            # Basic font setup
            run.font.size = Pt(11) # Standard script size
            # If it looks like a header (Short, all caps?), make bold?
            if len(full_text) < 50 and full_text.isupper() and align == WD_ALIGN_PARAGRAPH.CENTER:
                run.bold = True
            if "SCENE" in full_text.upper() or "LOCATION" in full_text.upper():
                run.bold = True

    # Add page number in footer
    section = doc.sections[-1]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = f"Page {page_num}"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page break after processing page (except last one handled by loop)
    # doc.add_page_break() # Handled in main loop

def pdf_to_docx(pdf_file, output_docx):
    print(f"Processing: {pdf_file}")
    
    # Step 1: Convert PDF to Images
    try:
        print("Converting PDF pages to images...")
        images = convert_from_path(pdf_file, poppler_path=POPPLER_PATH)
    except Exception as e:
        print(f"Error converting PDF to images: {e}")
        return

    doc = Document()
    
    print(f"Starting OCR and HOCR parsing for {len(images)} pages...")
    
    for i, image in enumerate(tqdm(images, desc="Processing Pages", unit="page")):
        # Preprocess
        processed_img = preprocess_image(image)
        
        try:
            # Get HOCR output
            hocr = pytesseract.image_to_pdf_or_hocr(
                processed_img, 
                extension='hocr',
                lang='eng+tam',
                config=TESSDATA_CONFIG + " -c hocr_font_info=1"
            )
            
            # Parse and write to DOCX
            hocr_to_docx(hocr, doc, i + 1)
            
            # Add page break between pages
            if i < len(images) - 1:
                doc.add_page_break()
                
        except pytesseract.TesseractError as e:
            print(f"Error on page {i+1}: {e}")
            doc.add_paragraph(f"[Error reading page {i+1}]")

    # Save
    doc.save(output_docx)
    print(f"Successfully saved to: {output_docx}")

def main():
    if len(sys.argv) > 1:
        input_path = sys.argv[1]
        if os.path.isdir(input_path):
            pdf_files = glob.glob(os.path.join(input_path, "*.pdf"))
            for pdf in pdf_files:
                out_name = os.path.splitext(pdf)[0] + ".docx"
                pdf_to_docx(pdf, out_name)
        elif os.path.isfile(input_path) and input_path.lower().endswith(".pdf"):
            out_name = os.path.splitext(input_path)[0] + ".docx"
            pdf_to_docx(input_path, out_name)
        else:
            print("Invalid input. Please provide a PDF file or directory.")
    else:
        print("Usage: python pdf_to_docx.py <path_to_pdf_or_directory>")
        path = input("Enter path to PDF file: ").strip().strip('"')
        if os.path.isfile(path):
            out_name = os.path.splitext(path)[0] + ".docx"
            pdf_to_docx(path, out_name)
        else:
            print("File not found.")

if __name__ == "__main__":
    main()
