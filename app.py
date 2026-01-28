import streamlit as st
import os
import sys
import numpy as np
import cv2
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
import io
import base64
from pypdf import PdfWriter, PdfReader

# ==============================================================================
# Configuration & Setup
# ==============================================================================

st.set_page_config(
    page_title="Phantom FX | Doc Tools",
    page_icon="photo header.png",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Tesseract Configuration
possible_tesseract_paths = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    os.path.join(os.getenv('LOCALAPPDATA'), r"Tesseract-OCR\tesseract.exe")
]
TESSERACT_CMD = None
for path in possible_tesseract_paths:
    if os.path.exists(path):
        TESSERACT_CMD = path
        break

if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# Poppler Configuration
# Check for local Windows path first, otherwise assume it's in system PATH (Linux/Cloud)
local_poppler_path = r"C:\Users\richardjoel.d\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe\poppler-25.07.0\Library\bin"

if os.path.exists(local_poppler_path):
    POPPLER_PATH = local_poppler_path
else:
    POPPLER_PATH = None  # pdf2image will look in system PATH

# Tesseract Data Path
TESSDATA_CONFIG = ""
if TESSERACT_CMD:
    tess_base = os.path.dirname(TESSERACT_CMD)
    sys_tessdata = os.path.join(tess_base, "tessdata")
    if not os.path.exists(os.path.join(sys_tessdata, "tam.traineddata")):
        local_tessdata = os.path.join(os.getcwd(), "tessdata")
        if os.path.exists(os.path.join(local_tessdata, "tam.traineddata")):
            os.environ["TESSDATA_PREFIX"] = local_tessdata

# ==============================================================================
# Assets & Helpers
# ==============================================================================

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

logo_path = "logo.png"
logo_base64 = ""
if os.path.exists(logo_path):
    logo_base64 = get_base64_of_bin_file(logo_path)

# Header image logic removed from body as per request


# ==============================================================================
# CSS & Styling (Glassmorphism + Phantom FX Theme)
# ==============================================================================

st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap');

    /* Global Reset & Body */
    .stApp {{
        background-color: #050505;
        background-image: 
            radial-gradient(circle at 50% 0%, rgba(47, 141, 77, 0.25), transparent 50%),
            radial-gradient(circle at 0% 50%, rgba(255, 20, 147, 0.1), transparent 50%);
        background-attachment: fixed;
        color: #f3f4f6;
        font-family: 'Outfit', sans-serif;
    }}
    
    header, footer, #MainMenu {{visibility: hidden;}}

    /* Main Container */
    .block-container {{
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 800px;
    }}

    /* Custom Header */
    .custom-header {{
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        gap: 1.5rem;
        margin-bottom: 3rem;
        padding: 1rem;
    }}
    
    .logo-container {{
        position: relative;
        padding: 1rem;
        border-radius: 1rem;
        background-color: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(8px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.2);
    }}

    .logo-img {{
        height: 56px;
        object-fit: contain;
    }}
    
    .header-title {{
        font-family: 'Outfit', sans-serif;
        font-weight: 700;
        font-size: 2.5rem;
        color: white;
        text-align: center;
        letter-spacing: -0.02em;
    }}
    
    .highlight-text {{
        color: #2F8D4D;
        text-shadow: 0 0 25px rgba(47, 141, 77, 0.5);
    }}

    /* Glass Card - Removed as per request to remove unwanted placeholders */
    .glass-card {{
        background-color: transparent;
        backdrop-filter: none;
        border: none;
        box-shadow: none;
        padding: 0;
        text-align: center;
    }}

    /* Upload Widget */
    [data-testid='stFileUploader'] {{
        padding: 2rem;
        border: 2px dashed rgba(255, 255, 255, 0.15);
        border-radius: 1rem;
        background-color: rgba(255, 255, 255, 0.02);
        transition: all 0.3s ease;
    }}
    [data-testid='stFileUploader']:hover {{
        border-color: rgba(52, 211, 153, 0.5);
        background-color: rgba(255, 255, 255, 0.05);
        box-shadow: 0 0 20px rgba(52, 211, 153, 0.15);
    }}

    /* Buttons */
    .stButton button {{
        background: linear-gradient(135deg, rgba(47, 141, 77, 0.1), rgba(47, 141, 77, 0.05));
        border: 1px solid rgba(47, 141, 77, 0.2);
        color: #2F8D4D;
        font-family: 'Outfit', sans-serif;
        font-weight: 600;
        border-radius: 0.75rem;
        padding: 0.75rem 2.5rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.2);
        width: 100%;
        margin-top: 1.5rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    .stButton button:hover {{
        background: linear-gradient(135deg, rgba(47, 141, 77, 0.2), rgba(47, 141, 77, 0.1));
        border-color: #2F8D4D;
        box-shadow: 0 0 25px rgba(47, 141, 77, 0.4);
        transform: translateY(-2px);
        color: white;
    }}

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {{
        gap: 1.5rem;
        background-color: transparent;
        justify-content: center;
        margin-bottom: 2rem;
    }}
    .stTabs [data-baseweb="tab"] {{
        background-color: rgba(255, 255, 255, 0.03);
        border-radius: 0.75rem;
        color: #9ca3af;
        border: 1px solid rgba(255, 255, 255, 0.1);
        padding: 0.75rem 1.5rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }}
    .stTabs [data-baseweb="tab"]:hover {{
        background-color: rgba(255, 255, 255, 0.08);
        color: white;
        border-color: rgba(255, 255, 255, 0.2);
    }}
    .stTabs [aria-selected="true"] {{
        background-color: rgba(47, 141, 77, 0.15) !important;
        border-color: rgba(47, 141, 77, 0.5) !important;
        color: white !important;
        box-shadow: 0 0 15px rgba(47, 141, 77, 0.2);
    }}

    /* Fix for Orange Underline - Make it match the theme */
    .stTabs [data-baseweb="tab-highlight"] {{
        background-color: #2F8D4D !important;
        box-shadow: 0 0 10px rgba(47, 141, 77, 0.5);
    }}

    /* Progress & Alerts */
    .stProgress > div > div > div > div {{
        background-color: #34d399;
        box-shadow: 0 0 15px rgba(52, 211, 153, 0.6);
    }}
    .stAlert {{
        background-color: rgba(52, 211, 153, 0.1);
        border: 1px solid rgba(52, 211, 153, 0.2);
        color: #d1fae5;
        backdrop-filter: blur(10px);
        border-radius: 0.75rem;
    }}

    h3 {{
        font-weight: 600;
        color: white !important;
        margin-bottom: 0.75rem;
        font-size: 1.5rem;
    }}
    p {{
        color: #9ca3af;
        font-weight: 400;
        font-size: 1rem;
        line-height: 1.6;
    }}
    .security-note {{
        margin-top: 4rem;
        text-align: center;
        font-size: 0.85rem;
        color: rgba(255, 255, 255, 0.3);
        letter-spacing: 0.05em;
    }}
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# Header
# ==============================================================================

st.markdown(f"""
<div class="custom-header">
    <div class="logo-container">
        <img src="data:image/png;base64,{logo_base64}" class="logo-img">
    </div>
    <div class="header-title">
        Phantom FX <span class="highlight-text">Doc Tools</span>
    </div>
</div>
""", unsafe_allow_html=True)

# ==============================================================================
# Logic Functions
# ==============================================================================

def preprocess_image(pil_image):
    open_cv_image = np.array(pil_image) 
    if len(open_cv_image.shape) == 3:
        open_cv_image = open_cv_image[:, :, ::-1].copy()

    gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(thresh)

def parse_bbox(title_str):
    if not title_str:
        return None
    match = re.search(r'bbox (\d+) (\d+) (\d+) (\d+)', title_str)
    if match:
        return [int(g) for g in match.groups()]
    return None

def hocr_to_docx(hocr_content, doc, page_num):
    """
    Robust HOCR parser that handles both PDF-based and Image-based HOCR outputs.
    Optimized for Script/Screenplay formatting (Tamil/English).
    """
    soup = BeautifulSoup(hocr_content, 'lxml')
    
    # 1. Determine Page Width
    page_width = 1000  # Default fallback
    page_div = soup.find('div', class_='ocr_page')
    if page_div:
        bbox = parse_bbox(page_div.get('title'))
        if bbox:
            page_width = bbox[2] - bbox[0]

    # 2. Extract Lines Directly (Stricter Line Preservation)
    # Instead of relying on paragraphs, we iterate lines to preserve script formatting exactly.
    lines = soup.find_all('span', class_='ocr_line')
    
    has_content = False
    
    for line in lines:
        # Check if line has actual text
        words = line.find_all('span', class_='ocrx_word')
        line_text_parts = []
        for word in words:
            text = word.get_text().strip()
            if text:
                line_text_parts.append(text)
        
        full_text = " ".join(line_text_parts).strip()
        
        if not full_text:
            continue
            
        docx_p = doc.add_paragraph()
        docx_p.paragraph_format.space_after = Pt(0) # Minimal spacing between lines to mimic PDF tight layout if needed, or Pt(6) for readabilty. 
        # For scripts, usually single spacing within blocks, double between blocks. 
        # But since we are mapping 1 line -> 1 para, let's keep it tight? 
        # User complained about "proper formatting", usually implies it looks like the PDF.
        # Let's use small space after.
        docx_p.paragraph_format.space_after = Pt(2)

        # Layout Analysis (Alignment/Indent)
        bbox = parse_bbox(line.get('title'))
        align = WD_ALIGN_PARAGRAPH.LEFT
        indent = 0
        
        if bbox:
            x0, y0, x1, y1 = bbox
            x_center = (x0 + x1) / 2
            page_center = page_width / 2
            
            # Script formatting heuristics
            # Center alignment check (looser tolerance)
            if abs(x_center - page_center) < (page_width * 0.15):
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif x1 > (page_width * 0.9) and x0 > (page_width * 0.6):
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif x0 > (page_width * 0.1):
                indent_ratio = x0 / page_width
                indent_inches = indent_ratio * 8.27
                indent = min(indent_inches, 4.0)

        docx_p.alignment = align
        if indent > 0:
            docx_p.paragraph_format.left_indent = Inches(indent)

        has_content = True
        
        # Script Dialogue Detection: "Name : Dialogue"
        # Pattern: Start of line, some text, spaces, colon, spaces, rest of text
        # We want to bold the "Name :" part
        dialogue_match = re.match(r'^([^:]+)(\s*:\s*)(.*)$', full_text)
        
        # Heuristic: Name shouldn't be too long (e.g. < 30 chars) to avoid false positives on regular sentences with colons
        is_dialogue = False
        if dialogue_match and len(dialogue_match.group(1)) < 30:
            is_dialogue = True
            name_part = dialogue_match.group(1)
            separator = dialogue_match.group(2)
            content_part = dialogue_match.group(3)
            
            # Add Name (Bold)
            run_name = docx_p.add_run(name_part)
            run_name.font.size = Pt(11)
            run_name.bold = True
            
            # Add Separator (Regular)
            run_sep = docx_p.add_run(separator)
            run_sep.font.size = Pt(11)
            
            # Add Content (Regular)
            run_content = docx_p.add_run(content_part)
            run_content.font.size = Pt(11)
            
        else:
            # Regular processing
            run = docx_p.add_run(full_text)
            run.font.size = Pt(11)
            
            # Basic styling heuristics for Headers/Scenes
            # 1. Short, Uppercase, Centered -> Likely Character Name (Standard format) or Title
            if len(full_text) < 50 and full_text.isupper() and align == WD_ALIGN_PARAGRAPH.CENTER:
                run.bold = True
            
            # 2. Explicit Scene Headings
            if any(keyword in full_text.upper() for keyword in ["SCENE:", "LOCATION:", "EFFECTS:", "காட்சி:", "இடம்:", "நேரம்:"]):
                run.bold = True
                
            # 3. Scene Summary Headers (Tamil)
            if "காட்சிச்சுருக்கம்" in full_text:
                run.bold = True


    # Fallback if HOCR failed to produce any text
    if not has_content:
        # Try raw text extraction if HOCR layout failed
        raw_text = soup.get_text()
        if raw_text.strip():
             doc.add_paragraph(raw_text.strip())
             has_content = True

    # Footer
    if page_num > 0:
        section = doc.sections[-1]
        footer = section.footer
        # Clear existing footer content if any (to avoid duplicates in loop)
        for p in footer.paragraphs:
            p.text = ""
        footer_p = footer.add_paragraph()
        footer_p.text = f"Page {page_num}"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return has_content

# ==============================================================================
# Main UI
# ==============================================================================

tab1, tab2, tab3 = st.tabs(["📄 PDF to Word", "🖼️ Image to Word", "📑 Merge PDFs"])

# ------------------------------------------------------------------------------
# Tab 1: PDF to Word
# ------------------------------------------------------------------------------
with tab1:
    st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
    st.markdown("<h3>Convert PDF to Word</h3>", unsafe_allow_html=True)
    st.markdown("<p>Convert scanned PDF scripts into editable Word documents.</p>", unsafe_allow_html=True)

    uploaded_pdf = st.file_uploader("Choose a PDF file", type="pdf", label_visibility="collapsed", key="pdf_uploader")

    if uploaded_pdf is not None:
        if st.button("Start Conversion", key="btn_pdf"):
            try:
                progress_bar = st.progress(0)
                status_text = st.empty()
                status_text.markdown("<p style='color: #34d399;'>Initializing...</p>", unsafe_allow_html=True)
                
                file_bytes = uploaded_pdf.read()
                images = convert_from_bytes(file_bytes, poppler_path=POPPLER_PATH)
                
                doc = Document()
                total_pages = len(images)
                
                for i, image in enumerate(images):
                    progress_bar.progress(int((i / total_pages) * 100))
                    status_text.markdown(f"<p style='color: #34d399;'>Processing Page {i+1} of {total_pages}...</p>", unsafe_allow_html=True)
                    
                    processed_img = preprocess_image(image)
                    
                    try:
                        hocr = pytesseract.image_to_pdf_or_hocr(
                            processed_img, 
                            extension='hocr',
                            lang='eng+tam',
                            config=TESSDATA_CONFIG + " -c hocr_font_info=1"
                        )
                        hocr_to_docx(hocr, doc, i + 1)
                        if i < total_pages - 1:
                            doc.add_page_break()
                    except Exception as e:
                        st.error(f"Error on page {i+1}: {e}")
                
                progress_bar.progress(100)
                status_text.markdown("<p style='color: #34d399;'>Conversion Complete!</p>", unsafe_allow_html=True)
                
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                st.success("✅ Document converted successfully!")
                st.download_button(
                    label="⬇️ Download Word Document",
                    data=docx_buffer,
                    file_name=f"{os.path.splitext(uploaded_pdf.name)[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"An error occurred: {e}")
    st.markdown("</div>", unsafe_allow_html=True)

# ------------------------------------------------------------------------------
# Tab 2: Image to Word
# ------------------------------------------------------------------------------
with tab2:
    st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
    st.markdown("<h3>Convert Image to Word</h3>", unsafe_allow_html=True)
    st.markdown("<p>Extract text from images (JPG/PNG) into editable Word documents.</p>", unsafe_allow_html=True)

    uploaded_img = st.file_uploader("Choose an Image", type=["jpg", "jpeg", "png"], label_visibility="collapsed", key="img_uploader")

    if uploaded_img is not None:
        if st.button("Start Conversion", key="btn_img"):
            try:
                progress_bar = st.progress(0)
                status_text = st.empty()
                status_text.markdown("<p style='color: #34d399;'>Processing Image...</p>", unsafe_allow_html=True)
                
                image = Image.open(uploaded_img)
                processed_img = preprocess_image(image)
                
                doc = Document()
                
                try:
                    # Get HOCR
                    hocr = pytesseract.image_to_pdf_or_hocr(
                        processed_img, 
                        extension='hocr',
                        lang='eng+tam',
                        config=TESSDATA_CONFIG + " -c hocr_font_info=1"
                    )
                    
                    # Debug: Check if HOCR is valid
                    if not hocr or len(hocr) < 10:
                        st.warning("OCR warning: Low confidence or empty output from Tesseract.")
                    
                    # Convert
                    has_content = hocr_to_docx(hocr, doc, 0) # 0 = no page number for single image
                    
                    if not has_content:
                         st.warning("No text could be detected in this image.")
                    else:
                        progress_bar.progress(100)
                        status_text.markdown("<p style='color: #34d399;'>Conversion Complete!</p>", unsafe_allow_html=True)
                        
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.success("✅ Image converted successfully!")
                        st.download_button(
                            label="⬇️ Download Word Document",
                            data=docx_buffer,
                            file_name=f"{os.path.splitext(uploaded_img.name)[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"OCR Logic Error: {e}")
                
            except Exception as e:
                st.error(f"An error occurred: {e}")
    st.markdown("</div>", unsafe_allow_html=True)

# ------------------------------------------------------------------------------
# Tab 3: Merge PDFs
# ------------------------------------------------------------------------------
with tab3:
    st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
    st.markdown("<h3>Merge Multiple PDFs</h3>", unsafe_allow_html=True)
    st.markdown("<p>Combine multiple PDF files into a single document.</p>", unsafe_allow_html=True)

    uploaded_pdfs = st.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True, label_visibility="collapsed", key="merge_uploader")

    if uploaded_pdfs:
        if st.button("Merge PDFs", key="btn_merge"):
            try:
                progress_bar = st.progress(0)
                status_text = st.empty()
                status_text.markdown("<p style='color: #34d399;'>Merging files...</p>", unsafe_allow_html=True)
                
                merger = PdfWriter()
                
                for i, pdf in enumerate(uploaded_pdfs):
                    # Create a PdfReader object for each uploaded file
                    # uploaded_pdfs are BytesIO-like objects
                    pdf_reader = PdfReader(pdf)
                    merger.append(pdf_reader)
                    progress_bar.progress(int((i + 1) / len(uploaded_pdfs) * 100))
                
                merged_buffer = io.BytesIO()
                merger.write(merged_buffer)
                merger.close()
                merged_buffer.seek(0)
                
                status_text.markdown("<p style='color: #34d399;'>Merge Complete!</p>", unsafe_allow_html=True)
                st.success("✅ PDFs merged successfully!")
                
                st.download_button(
                    label="⬇️ Download Merged PDF",
                    data=merged_buffer,
                    file_name="merged_document.pdf",
                    mime="application/pdf"
                )
                
            except Exception as e:
                st.error(f"Merge Error: {e}")
    st.markdown("</div>", unsafe_allow_html=True)

# Security Note
st.markdown("<div class='security-note'>🔒 All processing is done locally on this machine. No data is uploaded to external servers.</div>", unsafe_allow_html=True)
